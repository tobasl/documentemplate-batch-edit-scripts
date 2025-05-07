import os
from docx import Document

def replace_text_in_docx(file_path, search_text, replace_text):
    doc = Document(file_path)
    for para in doc.paragraphs:
        if search_text in para.text:
            para.text = para.text.replace(search_text, replace_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if search_text in cell.text:
                    cell.text = cell.text.replace(search_text, replace_text)

    doc.save(file_path)  # Overwrites the original file

def process_folder(folder_path, search_text, replace_text):
    for filename in os.listdir(folder_path):
        if filename.endswith(".docx"):
            file_path = os.path.join(folder_path, filename)
            print(f"Processing {filename}...")
            replace_text_in_docx(file_path, search_text, replace_text)

# === Customize these ===
folder = r"./"  # Change to your folder path
text_to_find = "<insert text>"
text_to_replace = "<insert text>"

process_folder(folder, text_to_find, text_to_replace)
